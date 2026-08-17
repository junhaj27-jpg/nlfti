from datetime import date,timedelta
from io import BytesIO
import pytest
from django.contrib.auth.models import Group,User
from django.core.exceptions import PermissionDenied,ValidationError
from docx import Document
from core.models import CAPA,Hazard,Nonconformity,Project,RiskAssessment,RiskControl
from core.quality import transition_capa
from core.reports import build_capa_report,build_risk_report

@pytest.fixture
def quality(db):
    analyst=User.objects.create_user("analyst"); reviewer=User.objects.create_user("reviewer"); rg=Group.objects.create(name="REVIEWER"); reviewer.groups.add(rg); project=Project.objects.create(title="Q",created_by=analyst)
    hazard=Hazard.objects.create(project=project,code="HZ-01",hazard="분할 누락",hazardous_situation="병변이 표시되지 않음")
    risk=RiskAssessment.objects.create(hazard=hazard,severity=5,probability=4,residual_severity=3,residual_probability=2,rationale="검토 통제 후 감소"); RiskControl.objects.create(assessment=risk,control_measure="이중 검토",implemented=True)
    nc=Nonconformity.objects.create(project=project,title="추론 실패",description="GPU OOM",source="ANALYSIS_FAILURE",created_by=analyst)
    capa=CAPA.objects.create(nonconformity=nc,owner=analyst,target_date=date.today()-timedelta(days=1),root_cause="ROI 과대",corrective_action="ROI 축소",preventive_action="메모리 사전 검사")
    return analyst,reviewer,project,risk,capa

@pytest.mark.django_db
def test_risk_rating_and_residual_risk(quality):
    risk=quality[3]; assert risk.initial_risk==20 and risk.initial_level=="HIGH"; assert risk.residual_risk==6 and risk.residual_level=="MEDIUM"
@pytest.mark.django_db
def test_capa_transition_permission_and_overdue(quality):
    analyst,reviewer,_,_,capa=quality; assert capa.overdue
    transition_capa(capa,"INVESTIGATING",analyst); transition_capa(capa,"ACTION_IN_PROGRESS",analyst); transition_capa(capa,"EFFECTIVENESS_REVIEW",analyst)
    with pytest.raises(PermissionDenied): transition_capa(capa,"CLOSED",analyst)
    capa.effectiveness_result="재실행 10건 이상 없음"; capa.save(); transition_capa(capa,"CLOSED",reviewer)
    assert capa.status=="CLOSED" and capa.closure_approved_by==reviewer and not capa.overdue
@pytest.mark.django_db
def test_invalid_capa_transition_rejected(quality):
    with pytest.raises(ValidationError): transition_capa(quality[4],"CLOSED",quality[1])
@pytest.mark.django_db
def test_quality_docx_reports(quality):
    project=quality[2]
    risk_text="\n".join(p.text for p in Document(BytesIO(build_risk_report(project))).paragraphs)
    capa_text="\n".join(p.text for p in Document(BytesIO(build_capa_report(project))).paragraphs)
    assert "잔여위험" in risk_text and "이중 검토" in risk_text
    assert "원인 분석" in capa_text and "효과성 검증" in capa_text
@pytest.mark.django_db
def test_dashboard_and_quality_pages_render(client,quality):
    analyst,_,project,_,_=quality; group=Group.objects.create(name="ANALYST"); analyst.groups.add(group); client.force_login(analyst)
    assert client.get("/").status_code==200
    assert client.get(f"/projects/{project.pk}/quality/").status_code==200
