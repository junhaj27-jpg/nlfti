from io import BytesIO
from docx import Document
from docx.shared import Inches

DISCLAIMER = "본 서비스는 연구·포트폴리오용이며 의료진의 진단을 대체하지 않습니다."
def build_ra_report(project):
    doc=Document(); doc.add_heading("의료영상 AI 분석·검증 보고서", 0); doc.add_paragraph(DISCLAIMER)
    doc.add_heading("1. 프로젝트 개요", 1); doc.add_paragraph(f"프로젝트명: {project.title}\n설명: {project.description or '-'}\n생성일: {project.created_at:%Y-%m-%d}")
    doc.add_heading("2. 입력 데이터 및 분석 결과", 1)
    for study in project.mri_studies.prefetch_related("jobs__reviews").all():
        doc.add_heading(f"BraTS MRI 4채널 / 검사일 {study.study_date}",2)
        doc.add_paragraph(f"입력 구성: T1, T1ce, T2, FLAIR{', 기준 마스크' if study.reference_mask else ''}\nShape: {study.shape}\nSpacing: {study.spacing} mm\n개인 식별정보는 수집하지 않음")
        doc.add_heading("전처리 과정",3); doc.add_paragraph("NIfTI 로드 → RAS orientation 통일 → 1 mm spacing 리샘플링 → 채널별 non-zero intensity 정규화 → foreground crop → 4채널 결합 → sliding-window 추론 → threshold → 작은 connected component 제거 → 원본 공간 복원")
        for job in study.jobs.all():
            doc.add_heading(f"분석 작업 #{job.pk}",3); doc.add_paragraph(f"사용 목적: 뇌종양 영역 분할 성능의 연구·포트폴리오 검증\n모델: {job.model_name} {job.model_version_name}\n모드/환경: {job.inference_mode} / {job.device or '-'} / mixed precision={job.mixed_precision}\n상태: {job.status}\n전처리 시간: {job.preprocessing_seconds or '-'} s\n추론 시간: {job.inference_seconds or '-'} s\n오류: {job.error_message or '-'}")
            if hasattr(job,"result"):
                r=job.result; doc.add_paragraph(f"전체 종양: {r.whole_tumor_cm3:.4f} cm³ ({r.whole_tumor_voxels} voxels)\n종양 핵심: {r.tumor_core_cm3:.4f} cm³ ({r.tumor_core_voxels} voxels)\n조영증강 종양: {r.enhancing_tumor_cm3:.4f} cm³ ({r.enhancing_tumor_voxels} voxels)")
                table=doc.add_table(rows=1,cols=5); table.style="Table Grid"
                for cell,text in zip(table.rows[0].cells,("영역","Dice","IoU","민감도","정밀도")): cell.text=text
                for region,metric in r.metrics.items():
                    cells=table.add_row().cells
                    for cell,text in zip(cells,(region,f"{metric['dice']:.4f}",f"{metric['iou']:.4f}",f"{metric['sensitivity']:.4f}",f"{metric['precision']:.4f}")): cell.text=text
                for correction in r.corrections.all(): doc.add_paragraph(f"마스크 수정본 #{correction.pk}: {correction.get_status_display()} / 수정자 {correction.editor.username} / 사유 {correction.reason} / 전체 종양 {correction.whole_tumor_cm3:.4f} cm³")
            doc.add_heading("검토·승인 이력",3)
            for review in job.reviews.all(): doc.add_paragraph(f"{review.reviewed_at:%Y-%m-%d %H:%M} / {review.reviewer.username} / {review.get_decision_display()} / {review.comment}")
    for image in project.images.prefetch_related("analyses__reviews").all():
        doc.add_heading(f"{image.modality} / 검사일 {image.study_date}", 2); doc.add_paragraph(f"익명화 파일 ID: {image.pk}\n설명: {image.description or '-'}")
        for a in image.analyses.all():
            doc.add_paragraph(f"모델: {a.model_version.name} {a.model_version.version}\n상태: {a.status}\n병변 voxel: {a.voxel_count}\nSpacing: {a.spacing_x}, {a.spacing_y}, {a.spacing_z} mm\n부피: {a.volume_cm3 if a.volume_cm3 is not None else '-'} cm³\n실행시간: {a.runtime_seconds if a.runtime_seconds is not None else '-'} s")
            doc.add_heading("성능평가", 3); table=doc.add_table(rows=1, cols=2); table.style="Table Grid"; table.rows[0].cells[0].text="지표"; table.rows[0].cells[1].text="값"
            for label,value in (("Dice = 2TP/(2TP+FP+FN)",a.dice),("IoU = TP/(TP+FP+FN)",a.iou),("민감도 = TP/(TP+FN)",a.sensitivity),("정밀도 = TP/(TP+FP)",a.precision)):
                cells=table.add_row().cells; cells[0].text=label; cells[1].text="-" if value is None else f"{value:.4f}"
            doc.add_heading("검토·승인 이력", 3)
            for r in a.reviews.all(): doc.add_paragraph(f"{r.reviewed_at:%Y-%m-%d %H:%M} / {r.reviewer.username} / {r.get_decision_display()} / {r.comment}")
    doc.add_heading("3. 실패 및 예외 기록",1)
    for study in project.mri_studies.all():
        for job in study.jobs.exclude(error_message=""): doc.add_paragraph(f"작업 #{job.pk}: {job.error_message}")
    doc.add_heading("4. 알려진 한계와 주의사항", 1); doc.add_paragraph("모델 성능은 사용한 가중치, BraTS 전처리 호환성, 입력 영상 품질에 의존합니다. mock 모드는 검증되지 않은 임계값 기반 결과입니다. 임상 사용, 진단, 치료 결정에 사용할 수 없으며 실제 의료기기 검증 및 규제 제출 요건을 충족하지 않습니다. 기준 마스크가 없는 경우 성능지표는 산출되지 않습니다.")
    stream=BytesIO(); doc.save(stream); return stream.getvalue()

def build_risk_report(project):
    doc=Document(); doc.add_heading("위험관리 요약 보고서",0); doc.add_paragraph(DISCLAIMER); doc.add_paragraph(f"프로젝트: {project.title}\n문서 목적: AI 의료영상 분석 관련 위험 식별, 평가, 통제 및 잔여위험 검토")
    for hazard in project.hazards.prefetch_related("assessments__controls").all():
        doc.add_heading(f"{hazard.code} · {hazard.hazard}",1); doc.add_paragraph(f"위해 상황: {hazard.hazardous_situation}\n위해: {hazard.harm or '-'}")
        table=doc.add_table(rows=1,cols=5); table.style="Table Grid"
        for cell,text in zip(table.rows[0].cells,("통제 전 S×P","통제 전 등급","통제 후 S×P","잔여위험 등급","근거")): cell.text=text
        for assessment in hazard.assessments.all():
            cells=table.add_row().cells
            for cell,text in zip(cells,(f"{assessment.severity}×{assessment.probability}={assessment.initial_risk}",assessment.initial_level,f"{assessment.residual_severity}×{assessment.residual_probability}={assessment.residual_risk}",assessment.residual_level,assessment.rationale)): cell.text=str(text)
            for control in assessment.controls.all(): doc.add_paragraph(f"통제: {control.control_measure} / 구현={control.implemented} / 검증={control.verification or '-'}")
    doc.add_heading("알려진 한계",1); doc.add_paragraph("위험등급은 포트폴리오용 5×5 행렬 예시이며 조직의 승인된 위험관리 절차와 ISO 14971 적용을 대체하지 않습니다.")
    out=BytesIO(); doc.save(out); return out.getvalue()

def build_capa_report(project):
    doc=Document(); doc.add_heading("부적합 및 CAPA 보고서",0); doc.add_paragraph(DISCLAIMER); doc.add_paragraph(f"프로젝트: {project.title}")
    for nc in project.nonconformities.select_related("capa").all():
        doc.add_heading(f"NC-{nc.pk}: {nc.title}",1); doc.add_paragraph(f"출처: {nc.get_source_display()}\n부적합 내용: {nc.description}\n생성일: {nc.created_at:%Y-%m-%d}")
        if hasattr(nc,"capa"):
            capa=nc.capa; doc.add_heading(f"CAPA-{capa.pk}",2); doc.add_paragraph(f"상태: {capa.get_status_display()}\n원인 분석: {capa.root_cause or '-'}\n시정조치: {capa.corrective_action or '-'}\n예방조치: {capa.preventive_action or '-'}\n담당자: {capa.owner.username}\n목표일: {capa.target_date}\n기한 초과: {capa.overdue}\n효과성 검증: {capa.effectiveness_result or '-'}\n종료 승인자: {capa.closure_approved_by.username if capa.closure_approved_by else '-'}")
        else: doc.add_paragraph("연결된 CAPA가 없습니다.")
    out=BytesIO(); doc.save(out); return out.getvalue()
